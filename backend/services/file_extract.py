"""여러 앱의 "붙여넣기" 텍스트 입력에 공통으로 붙는 파일 업로드 기능 — Word/PDF/Excel처럼
텍스트가 아닌 바이너리 형식에서 순수 텍스트를 뽑아내 기존 textarea에 채워 넣을 수 있게 한다.

txt/csv/json/log 등 이미 텍스트인 형식은 브라우저에서 FileReader로 바로 읽어도 되지만,
Word/PDF/Excel은 실제 파싱이 필요해 서버에서 처리한다 — 그래서 프론트는 파일 형식과
무관하게 항상 `/api/extract-text`로 보내고, 이 함수가 확장자별로 적절히 분기한다.

⚠️ 원본 파일은 어디에도 저장하지 않고 추출된 텍스트만 반환한다 — App 19(시크릿 스캐너)의
"원본 미저장" 원칙과 동일.
"""
import csv
import io

MAX_CHARS = 100_000

_TEXT_EXTENSIONS = {
    "txt", "log", "md", "json", "yml", "yaml", "conf", "cfg", "ini", "xml", "html", "py", "js",
    "sh", "ps1", "toml", "env",
}


class ExtractError(ValueError):
    pass


def _extract_txt(raw: bytes) -> str:
    return raw.decode("utf-8", errors="replace")


def _extract_csv(raw: bytes) -> str:
    # CSV도 사실 텍스트지만, 구분자가 깨져 보이지 않도록 한 번 파싱해 정렬된 형태로 되돌린다.
    text = raw.decode("utf-8-sig", errors="replace")
    try:
        reader = csv.reader(io.StringIO(text))
        rows = [", ".join(row) for row in reader]
        return "\n".join(rows)
    except Exception:
        return text


def _extract_docx(raw: bytes) -> str:
    try:
        import docx
    except ImportError as e:
        raise ExtractError("서버에 python-docx가 설치되어 있지 않습니다.") from e

    doc = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    if not parts:
        raise ExtractError("문서에서 텍스트를 찾지 못했습니다 (이미지만 있는 문서일 수 있습니다).")
    return "\n".join(parts)


def _extract_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ExtractError("서버에 pypdf가 설치되어 있지 않습니다.") from e

    reader = PdfReader(io.BytesIO(raw))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as e:
            raise ExtractError("암호로 보호된 PDF는 지원하지 않습니다. 암호를 해제한 뒤 업로드하세요.") from e

    parts = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)
    if not parts:
        raise ExtractError("PDF에서 텍스트를 찾지 못했습니다 (스캔 이미지 PDF일 수 있습니다 — OCR이 필요합니다).")
    return "\n\n".join(parts)


def _extract_xlsx(raw: bytes) -> str:
    try:
        import openpyxl
    except ImportError as e:
        raise ExtractError("서버에 openpyxl이 설치되어 있지 않습니다.") from e

    try:
        wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    except Exception as e:
        raise ExtractError(f"엑셀 파일을 읽지 못했습니다: {e}") from e

    parts = []
    for sheet in wb.worksheets:
        parts.append(f"--- 시트: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append(", ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        raise ExtractError("엑셀 파일에서 데이터를 찾지 못했습니다.")
    return text


_EXTRACTORS = {
    "docx": _extract_docx,
    "pdf": _extract_pdf,
    "xlsx": _extract_xlsx,
    "xls": _extract_xlsx,
    "csv": _extract_csv,
}


def extract_text(filename: str, raw: bytes) -> dict:
    """파일명·바이트로부터 순수 텍스트를 추출한다. 반환: {"text", "truncated"}."""
    if not raw:
        raise ExtractError("빈 파일입니다.")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    extractor = _EXTRACTORS.get(ext)

    if extractor:
        text = extractor(raw)
    elif ext in _TEXT_EXTENSIONS or not ext:
        text = _extract_txt(raw)
    else:
        # 모르는 확장자는 일단 텍스트로 시도 — 바이너리면 대부분 깨진 문자로 나타나 사용자가 바로 알아챌 수 있음.
        text = _extract_txt(raw)

    truncated = len(text) > MAX_CHARS
    if truncated:
        text = text[:MAX_CHARS]
    return {"text": text, "truncated": truncated}
